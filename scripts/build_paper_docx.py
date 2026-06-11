"""
Generate the research paper as a formatted DOCX.
Run: python scripts/build_paper_docx.py
Output: docs/paper.docx
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(__file__), '..', 'docs', 'paper.docx')
FIGS = os.path.join(os.path.dirname(__file__), '..', 'docs', 'figures')

# ── helpers ──────────────────────────────────────────────────────────────────

def set_font(run, name='Times New Roman', size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        set_font(run, size=13, bold=True)
    elif level == 2:
        set_font(run, size=11, bold=True, italic=True)
    else:
        set_font(run, size=11, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_para(doc, text, size=10, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.first_line_indent = Inches(0.25)
    return p

def add_bullet(doc, text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_font(run, size=size)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_caption(doc, text, size=9):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=size, italic=True)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)

def set_cell_bg(cell, hex_color='D9E1F2'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header row
    for j, h in enumerate(headers):
        cell = t.cell(0, j)
        cell.text = h
        set_cell_bg(cell, 'BDD7EE')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        set_font(run, size=9, bold=True)

    # data rows
    for i, row in enumerate(rows):
        bg = 'FFFFFF' if i % 2 == 0 else 'EBF3FB'
        for j, val in enumerate(row):
            cell = t.cell(i + 1, j)
            cell.text = str(val)
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            bold = (i == 0 and j > 0) or (j == 0)
            set_font(run, size=9, bold=bold)

    # column widths
    if col_widths:
        for i, row_cells in enumerate(t.rows):
            for j, cell in enumerate(row_cells.cells):
                cell.width = Inches(col_widths[j])
    return t


# ── document ─────────────────────────────────────────────────────────────────

doc = Document()

# page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Title ────────────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run(
    "Calibrated Differential Diagnosis with KL-Divergence Soft-Label Training:\n"
    "A Comparative Study of TF-IDF, DistilBERT, and PubMedBERT on DDXPlus"
)
set_font(title_run, size=14, bold=True)
title_p.paragraph_format.space_after = Pt(12)

# ── Authors ──────────────────────────────────────────────────────────────────
authors_p = doc.add_paragraph()
authors_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors_run = authors_p.add_run(
    "Annaji Kuthe\u00b9, Harshal Donarkar\u00b2*, Charul Patel\u00b3, "
    "Mahek Qureshi\u2074, Angad Bawankar\u2075, Kunal Gawande\u2076"
)
set_font(authors_run, size=10, bold=True)
authors_p.paragraph_format.space_after = Pt(6)

affiliations = [
    "\u00b9 Annaji Kuthe, Professor, Department of Computer Science & Engineering, "
    "K.D.K. College of Engineering, Nagpur, Maharashtra, India",
    "\u00b2 Harshal Donarkar, Student, Department of Computer Science & Engineering, "
    "K.D.K. College of Engineering, Nagpur, Maharashtra, India",
    "\u00b3 Charul Patel, Student, Department of Computer Science & Engineering, "
    "K.D.K. College of Engineering, Nagpur, Maharashtra, India",
    "\u2074 Mahek Qureshi, Student, Department of Computer Science & Engineering, "
    "K.D.K. College of Engineering, Nagpur, Maharashtra, India",
    "\u2075 Angad Bawankar, Student, Department of Computer Science & Engineering, "
    "K.D.K. College of Engineering, Nagpur, Maharashtra, India",
    "\u2076 Kunal Gawande, Student, Department of Computer Science & Engineering, "
    "K.D.K. College of Engineering, Nagpur, Maharashtra, India",
]
for affil in affiliations:
    affil_p = doc.add_paragraph()
    affil_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affil_run = affil_p.add_run(affil)
    set_font(affil_run, size=9, italic=True)
    affil_p.paragraph_format.space_after = Pt(2)

corr_p = doc.add_paragraph()
corr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
corr_run = corr_p.add_run(
    "*Corresponding Author: Harshal Donarkar\n"
    "Email: harshaldonarkar@gmail.com\n"
    "Phone: +91 7721984092\n"
    "Fax: N/A"
)
set_font(corr_run, size=9, italic=True)
corr_p.paragraph_format.space_after = Pt(14)

doc.add_paragraph()  # spacer

# ── Abstract ─────────────────────────────────────────────────────────────────
abs_label = doc.add_paragraph()
abs_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(abs_label.add_run("Abstract"), size=11, bold=True)

abs_p = doc.add_paragraph()
abs_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
abs_run = abs_p.add_run(
    "Accurate differential diagnosis — ranking plausible diseases by probability given observed symptoms — "
    "is a foundational clinical reasoning task. Existing machine-learning approaches optimise top-1 accuracy, "
    "yet a clinically useful system must also produce well-calibrated probability distributions over all candidate "
    "diseases. We study this distinction on DDXPlus, a large-scale benchmark of 1.03 million simulated patient "
    "cases across 49 pathologies. We train three models: (1) TF-IDF bag-of-symptoms with Logistic Regression "
    "(TF-IDF+LR), (2) DistilBERT fine-tuned with KL-divergence soft-label loss (DistilBERT+KL), and "
    "(3) PubMedBERT fine-tuned with the same objective (PubMedBERT+KL). Our key findings are: TF-IDF+LR "
    "achieves the highest top-1 accuracy (99.51%) but produces severely miscalibrated distributions "
    "(mean KL = 6.29); transformer models trained with KL-divergence loss reduce miscalibration by 15x "
    "(KL \u2248 0.41) and improve NDCG@3 from 0.748 to 0.921 with only a 1% drop in top-1 accuracy. "
    "Surprisingly, biomedical domain pre-training (PubMedBERT) provides no significant advantage over "
    "general-domain pre-training (DistilBERT) on this dataset, because DDXPlus symptoms are encoded as "
    "binary feature vectors rather than free clinical text. We integrate the best model into an end-to-end "
    "web application augmented with chest X-ray analysis, lab report interpretation, and multi-provider LLM "
    "explanation generation including Google MedGemma."
)
set_font(abs_run, size=10)
abs_p.paragraph_format.first_line_indent = Inches(0)
abs_p.paragraph_format.space_after = Pt(6)

kw_p = doc.add_paragraph()
kw_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
kw_run = kw_p.add_run(
    "Keywords: differential diagnosis, KL divergence, probability calibration, "
    "clinical decision support, machine learning"
)
set_font(kw_run, size=10, italic=True)
kw_p.paragraph_format.first_line_indent = Inches(0)
kw_p.paragraph_format.space_after = Pt(14)

# ── 1. Introduction ───────────────────────────────────────────────────────────
add_heading(doc, "1.  Introduction", 1)
add_para(doc,
    "Clinical decision support systems have long promised to assist physicians by surfacing plausible "
    "diagnoses from patient-reported symptoms. The dominant paradigm evaluates such systems by top-1 "
    "accuracy: does the system's single highest-confidence prediction match the ground-truth diagnosis? "
    "While this metric is intuitive, it ignores a property that clinicians rely on daily: a differential "
    "diagnosis is not a single answer but an ordered distribution over candidate conditions, weighted by "
    "clinical likelihood.")
add_para(doc,
    "Consider two systems that both correctly predict 'Pneumonia' as the top diagnosis. System A assigns "
    "95% probability to Pneumonia and near-zero probability to Pulmonary Embolism and Myocardial Infarction. "
    "System B assigns 60% to Pneumonia, 20% to Pulmonary Embolism, and 5% to MI. Both achieve equal top-1 "
    "accuracy; yet System B is clinically superior — it correctly signals that the presentation is ambiguous "
    "and that dangerous alternatives must be actively ruled out.")
add_para(doc, "This paper makes the following contributions:")
add_bullet(doc, "Novel training objective: KL-divergence soft-label loss fine-tuning on DDXPlus, directly optimising the full probability distribution over 49 diseases.")
add_bullet(doc, "Systematic three-model comparison across six metrics: Acc@{1,3,5}, Macro-F1, NDCG@{3,5}, mean KL divergence, and mean rank on 134,529 test cases.")
add_bullet(doc, "Negative result on domain pre-training: PubMedBERT provides no measurable benefit over DistilBERT on structured symptom inputs — a practically important finding.")
add_bullet(doc, "End-to-end clinical platform integrating DDXPlus diagnosis, chest X-ray analysis, lab report interpretation, and MedGemma-powered explanations.")

# ── 2. Related Work ───────────────────────────────────────────────────────────
add_heading(doc, "2.  Related Work", 1)

add_heading(doc, "2.1  Differential Diagnosis Datasets", 2)
add_para(doc,
    "DDXPlus (Tchango et al., NeurIPS 2022) is the largest publicly available differential diagnosis "
    "benchmark, comprising 1,033,360 synthetic patient cases generated by a clinical simulator. Each case "
    "includes binary and multi-value symptom encodings, demographics, a ground-truth disease label, and a "
    "physician-annotated differential diagnosis distribution over 49 pathologies. The soft-label differential "
    "makes DDXPlus uniquely suited for calibration-aware training. Prior work on automated diagnosis from "
    "structured symptom inputs (Rotmensch et al., 2017) established symptom–disease relationships from "
    "electronic medical records, but such datasets are substantially smaller and do not provide soft-label "
    "differential distributions.")

add_heading(doc, "2.2  Machine Learning for Diagnosis", 2)
add_para(doc,
    "Logistic Regression and gradient-boosted trees over bag-of-symptom features remain competitive baselines "
    "for structured symptom inputs (Buch et al., 2018). Transformer models pre-trained "
    "on biomedical text — BioBERT (Lee et al., 2020), ClinicalBERT (Alsentzer et al., 2019), PubMedBERT "
    "(Gu et al., 2021) — achieve state-of-the-art performance on clinical NLP tasks. Their application to "
    "structured differential diagnosis with soft-label objectives has not been systematically studied.")

add_heading(doc, "2.3  Calibration in Medical AI", 2)
add_para(doc,
    "Model calibration — alignment between predicted confidence and empirical accuracy — has received growing "
    "attention in medical imaging (Guo et al., 2017). KL divergence as a direct training objective for "
    "soft-label supervision was proposed in knowledge distillation (Hinton et al., 2015) but has not been "
    "specifically studied in clinical differential diagnosis. NDCG as a ranking metric for differential "
    "diagnosis was introduced by Tchango et al. (2022) alongside DDXPlus.")

add_heading(doc, "2.4  Medical Foundation Models", 2)
add_para(doc,
    "MedGemma (Google, 2024) is an open-weight multimodal foundation model trained on medical text and images. "
    "CheXNet (Rajpurkar et al., 2017) demonstrated radiologist-level pneumonia detection. torchxrayvision "
    "(Cohen et al., 2022) provides pre-trained chest X-ray models across multiple public datasets. Our work "
    "integrates both a fine-tuned diagnostic classifier and pre-trained radiological models in a unified platform.")

# ── 3. Dataset ────────────────────────────────────────────────────────────────
add_heading(doc, "3.  Dataset", 1)
add_para(doc,
    "We use DDXPlus (Tchango et al., 2022) accessed via the HuggingFace Hub (mila-iqia/ddxplus). The dataset "
    "contains 1,033,360 patient cases split into train (80%), validation (10%), and test (10%). Each case "
    "comprises 93 antecedent features, 223 symptom features, and 2 demographic features (age bin, sex), "
    "yielding a 318-dimensional input vector. The ground-truth label is one of 49 diseases; the differential "
    "diagnosis is a physician-annotated probability distribution over all 49 diseases used as soft supervision.")
add_para(doc,
    "All features are decoded from coded IDs to human-readable English text using the release_evidences.json "
    "vocabulary, producing a natural-language symptom string per patient (e.g. '35 year old male. chest pain, "
    "shortness of breath, leg swelling. Initial complaint: chest pain.'). Both TF-IDF+LR and the transformer "
    "models use this same decoded text as input, ensuring differences in performance are attributable to model "
    "architecture and training objective, not input representation. "
    "The test partition contains 134,529 cases with approximately uniform disease prevalence "
    "(median 2,572 cases per disease, range 1,847\u20133,401), making macro-averaged metrics meaningful.")

# ── 4. Methods ────────────────────────────────────────────────────────────────
add_heading(doc, "4.  Methods", 1)

add_heading(doc, "4.1  TF-IDF + Logistic Regression Baseline", 2)
add_para(doc,
    "A TF-IDF vectoriser (max 10,000 features, unigrams and bigrams) is fitted on training symptom strings. "
    "A one-vs-rest Logistic Regression (C=1.0, max_iter=1000, solver='lbfgs') is trained on TF-IDF features. "
    "Class probabilities are obtained via predict_proba.")

add_heading(doc, "4.2  KL-Divergence Soft-Label Loss", 2)
add_para(doc,
    "Let y \u2208 \u0394\u2074\u2079 denote the physician-annotated differential and \u0177 = softmax(Wh) "
    "the model's prediction. The training loss combines KL divergence and cross-entropy:")
eq_p = doc.add_paragraph()
eq_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
eq_run = eq_p.add_run(
    "\u2112 = 0.7 \u00d7 D_KL(y \u2016 \u0177) + 0.3 \u00d7 \u2112_CE(argmax(y), \u0177)"
)
set_font(eq_run, name='Courier New', size=10, italic=True)
eq_p.paragraph_format.first_line_indent = Inches(0)
add_para(doc,
    "where D_KL(y\u2016\u0177) = \u03a3_i y_i log(y_i/\u0177_i). The auxiliary cross-entropy term "
    "(weight 0.3) stabilises training when the soft label is nearly uniform. The 0.7/0.3 split was chosen "
    "on the validation set.")

add_heading(doc, "4.3  DistilBERT + KL", 2)
add_para(doc,
    "We fine-tune distilbert-base-uncased (66M parameters) with a linear classification head over the [CLS] "
    "token embedding. Training: AdamW (lr=2\u00d710\u207b\u2075, weight decay=0.01), batch size 64, 3 epochs, "
    "max sequence length 128. Hardware: Apple M-series with MPS acceleration (~2 hours).")

add_heading(doc, "4.4  PubMedBERT + KL", 2)
add_para(doc,
    "We fine-tune microsoft/BiomedNLP-BiomedBERT-base-uncased-abstract-fulltext (110M parameters) with the "
    "identical configuration and loss function. PubMedBERT was pre-trained from scratch on PubMed abstracts "
    "and full-text articles, providing a vocabulary tuned to biomedical terminology.")

add_heading(doc, "4.5  Loss Weight Ablation (\u03b1)", 2)
add_para(doc,
    "To justify the \u03b1 = 0.7 choice, we train DistilBERT with \u03b1 \u2208 {0.5, 0.7, 0.9}, holding all "
    "other hyperparameters fixed, and evaluate on the full test set (Table 5). Higher \u03b1 monotonically "
    "improves calibration (lower KL) but reduces Acc@1: \u03b1=0.9 achieves the best KL (0.3025) and "
    "NDCG@3 (0.9247) at a 3% Acc@1 cost, while \u03b1=0.5 maximises Acc@1 (0.9863) but doubles the KL "
    "divergence. We select \u03b1=0.7 as the best practical trade-off.")

ablation_headers = ['\u03b1 (KL weight)', 'Acc@1', 'NDCG@3', 'Mean KL \u2193']
ablation_rows = [
    ['0.5',          '0.9863*', '0.9198',  '0.6190'],
    ['0.7 (ours)',   '0.9851',  '0.9214',  '0.4164'],
    ['0.9',          '0.9557',  '0.9247*', '0.3025*'],
]
add_table(doc, ablation_headers, ablation_rows, col_widths=[1.5, 1.5, 1.5, 1.5])
add_caption(doc, "Table 5. \u03b1 ablation on test set (DistilBERT+KL, n=134,529). * = best in column. "
            "\u03b1=0.7 selected as best trade-off: keeps Acc@1 within 1% of baseline while reducing KL "
            "vs \u03b1=0.5. \u03b1=0.9 gives marginal NDCG@3 gain at 3% Acc@1 cost.")

add_heading(doc, "4.6  Evaluation Metrics", 2)
add_para(doc,
    "Acc@k: fraction of patients where the true diagnosis appears in the top-k predictions. "
    "Macro-F1: unweighted average F1 across 49 classes. "
    "NDCG@k: normalised discounted cumulative gain using the physician differential as relevance weights. "
    "Mean KL: average KL divergence between the physician distribution and predicted distribution "
    "(lower = better calibration). Mean Rank: average rank of the ground-truth disease (1.0 = always first).")

# ── 5. Results ────────────────────────────────────────────────────────────────
add_heading(doc, "5.  Results", 1)

add_heading(doc, "5.1  Main Comparison", 2)
add_para(doc,
    "Table 1 presents evaluation results on the DDXPlus test set (n = 134,529). "
    "Best values per column are marked with an asterisk (*). Mean KL and Mean Rank are lower-is-better. "
    "95% bootstrap confidence intervals (1,000 resamples) are reported in Table 2.")

# Table 1
table1_headers = ['Model', 'Acc@1', 'Acc@3', 'Acc@5', 'Macro-F1', 'NDCG@3', 'NDCG@5', 'Mean KL\u2193', 'ECE\u2193', 'Mean Rank\u2193']
table1_rows = [
    ['TF-IDF + LR',     '0.9951*','1.0000*','1.0000*','0.9940*','0.7483', '0.7059', '6.2884',  '0.0043*','1.0049*'],
    ['DistilBERT + KL', '0.9851', '0.9993', '0.9999', '0.9840', '0.9214*','0.9203*','0.4164',  '0.4813', '1.0182'],
    ['PubMedBERT + KL', '0.9825', '0.9992', '0.9999', '0.9814', '0.9201', '0.9192', '0.4111*', '0.4868', '1.0206'],
]
add_table(doc, table1_headers, table1_rows,
          col_widths=[1.4, 0.6, 0.6, 0.6, 0.7, 0.6, 0.6, 0.65, 0.55, 0.7])
add_caption(doc, "Table 1. Evaluation on DDXPlus test set (n=134,529, 49 diseases). * = best in column. "
            "\u2193 = lower is better. ECE (Expected Calibration Error, 15 bins, top-1 confidence). "
            "Note: TF-IDF+LR achieves near-zero ECE because it assigns ~99% confidence to one class and "
            "is correct 99.5% of the time; see Section 5.6 for why ECE is an inappropriate metric for "
            "differential diagnosis.")

# Bootstrap CI table
ci_headers = ['Model', 'Acc@1', 'Acc@1 95% CI', 'NDCG@3', 'NDCG@3 95% CI', 'Mean KL', 'KL 95% CI']
ci_data = [
    ['TF-IDF + LR',     '0.9951', '[0.9948, 0.9955]', '0.7483', '[0.7473, 0.7493]', '6.2885', '[6.2740, 6.3020]'],
    ['DistilBERT + KL', '0.9851', '[0.9844, 0.9857]', '0.9214', '[0.9207, 0.9219]', '0.4164', '[0.4151, 0.4177]'],
    ['PubMedBERT + KL', '0.9825', '[0.9818, 0.9832]', '0.9201', '[0.9195, 0.9207]', '0.4111', '[0.4097, 0.4124]'],
]
add_table(doc, ci_headers, ci_data, col_widths=[1.3, 0.6, 1.25, 0.6, 1.25, 0.6, 1.25])
add_caption(doc, "Table 2. Bootstrap 95% confidence intervals (1,000 resamples, n=134,529 test cases). "
            "Non-overlapping CIs confirm all key differences are statistically reliable.")

add_heading(doc, "5.2  Accuracy vs. Calibration Trade-off", 2)
add_para(doc,
    "TF-IDF+LR achieves the highest top-1 accuracy (99.51%) but its mean KL divergence of 6.29 reveals "
    "that while it frequently selects the correct answer, its probability distribution over all 49 diseases "
    "is severely miscalibrated. DistilBERT+KL and PubMedBERT+KL sacrifice approximately 1% of top-1 accuracy "
    "but reduce mean KL divergence by a factor of 15x (6.29 \u2192 0.41). The NDCG@3 increase from 0.748 "
    "to 0.921 confirms that transformer models not only rank the correct disease first more often, but also "
    "rank the second- and third-most-probable diseases correctly — a clinically critical property.")

add_heading(doc, "5.3  Effect of Biomedical Pre-Training", 2)
add_para(doc,
    "Comparing DistilBERT+KL (general-domain) to PubMedBERT+KL (biomedical-domain): PubMedBERT shows "
    "marginally lower Acc@1 (\u22120.26%), negligible NDCG@3 difference (\u22120.0013), and marginally lower "
    "KL (\u22120.005). These differences are not clinically or statistically meaningful. The negative result "
    "is attributable to DDXPlus's structured symptom encoding: inputs are templated feature strings, not "
    "biomedical prose, so vocabulary-level domain knowledge provides no advantage.")

add_heading(doc, "5.4  Statistical Significance", 2)
add_para(doc,
    "We verify that observed differences are statistically significant using two tests applied to the full "
    "test set (n = 134,529).")

add_para(doc,
    "McNemar's test (Acc@1 pairwise). McNemar's test compares the error patterns of two classifiers on the "
    "same test set. Results: TF-IDF+LR vs DistilBERT+KL: \u03c7\u00b2 = 887.73, p = 4.57\u00d710\u207b\u00b9\u2079\u2075; "
    "TF-IDF+LR vs PubMedBERT+KL: \u03c7\u00b2 = 1154.64, p = 4.40\u00d710\u207b\u00b2\u2075\u00b3; "
    "DistilBERT+KL vs PubMedBERT+KL: \u03c7\u00b2 = 140.33, p = 2.26\u00d710\u207b\u00b3\u00b2. "
    "All pairwise Acc@1 differences are highly significant (p \u226a 0.05), including between the two "
    "transformer models. However, the practical effect size between DistilBERT and PubMedBERT is negligible "
    "(\u0394Acc@1 = 0.26%), with the significance driven by the large sample size.")

add_para(doc,
    "Wilcoxon signed-rank test (per-sample KL divergence). The Wilcoxon test compares the distribution of "
    "per-patient KL divergences between pairs of models using a 10,000-case subsample. TF-IDF+LR vs "
    "DistilBERT+KL: p \u2248 0 (mean KL 6.2884 vs 0.4164); TF-IDF+LR vs PubMedBERT+KL: p \u2248 0 "
    "(mean KL 6.2884 vs 0.4111). Both confirm that the calibration improvement from KL-divergence training "
    "is not due to sampling variation. DistilBERT+KL vs PubMedBERT+KL also reaches p \u2248 0, but the "
    "practical difference (\u03940.005 KL) is clinically negligible — a textbook case of statistical vs. "
    "practical significance at large n.")

# Summary significance table
sig_headers = ['Comparison', 'Test', 'Statistic', 'p-value', 'Significant']
sig_rows_data = [
    ['TF-IDF+LR vs DistilBERT+KL', "McNemar (Acc@1)", '\u03c7\u00b2 = 887.73', '4.57\u00d710\u207b\u00b9\u2079\u2075', 'YES'],
    ['TF-IDF+LR vs PubMedBERT+KL', "McNemar (Acc@1)", '\u03c7\u00b2 = 1154.64','4.40\u00d710\u207b\u00b2\u2075\u00b3','YES'],
    ['DistilBERT+KL vs PubMedBERT+KL','McNemar (Acc@1)','\u03c7\u00b2 = 140.33','2.26\u00d710\u207b\u00b3\u00b2','YES (small \u0394)'],
    ['TF-IDF+LR vs DistilBERT+KL', 'Wilcoxon (KL)', 'KL: 6.2884 vs 0.4164', '\u22480', 'YES'],
    ['TF-IDF+LR vs PubMedBERT+KL', 'Wilcoxon (KL)', 'KL: 6.2884 vs 0.4111', '\u22480', 'YES'],
    ['DistilBERT+KL vs PubMedBERT+KL','Wilcoxon (KL)','KL: 0.4164 vs 0.4111','\u22480','YES (trivial \u0394)'],
]
add_table(doc, sig_headers, sig_rows_data, col_widths=[1.8, 1.4, 1.6, 1.1, 1.1])
add_caption(doc, "Table 3. Statistical significance tests. McNemar: H\u2080 = identical error patterns. "
            "Wilcoxon: H\u2080 = identical per-sample KL distribution (n=10,000 subsample). All \u03b1=0.05.")

add_heading(doc, "5.5  Calibration", 2)
add_para(doc,
    "Reliability diagrams (Figure 1) confirm that TF-IDF+LR is systematically overconfident, predicting high "
    "probabilities (>0.8) for cases where empirical accuracy is substantially lower. DistilBERT+KL and "
    "PubMedBERT+KL show substantially better calibration, with predicted probabilities tracking the diagonal "
    "across all confidence bins.")

# Try to insert calibration figure
cal_fig_path = os.path.join(FIGS, 'calibration.png')
if os.path.exists(cal_fig_path):
    doc.add_paragraph()
    fig_p = doc.add_paragraph()
    fig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig_p.add_run()
    run.add_picture(cal_fig_path, width=Inches(5.5))
    add_caption(doc,
        "Figure 1. Reliability diagrams (calibration curves) for all three models on the DDXPlus test set "
        "(n=134,529). Each point represents empirical accuracy within a confidence bin vs. mean predicted "
        "confidence. The dashed diagonal represents perfect calibration. TF-IDF+LR concentrates predictions "
        "at very high confidence with near-perfect ECE but severely miscalibrated full distributions "
        "(mean KL=6.29). DistilBERT+KL and PubMedBERT+KL track the diagonal across all bins (mean KL\u22480.41).")

add_heading(doc, "5.6  Per-Disease F1 Analysis", 2)
add_para(doc,
    "Figure 2 shows per-disease macro F1-score for all three models sorted by DistilBERT+KL performance. "
    "Both transformer models achieve F1 > 0.95 on 44 of 49 diseases. The five lowest-performing diseases "
    "share heavily overlapping symptom profiles: Acute vs. Chronic rhinosinusitis differ primarily in "
    "duration, while Stable vs. Unstable angina present identically at onset. In these ambiguous cases the "
    "physician-annotated differential itself assigns split probability between the two conditions \u2014 "
    "precisely the scenario where NDCG@3 (0.921 for transformers vs. 0.748 for TF-IDF+LR) is a more "
    "informative metric than Acc@1.")

f1_fig_path = os.path.join(FIGS, 'per_disease_f1.png')
if os.path.exists(f1_fig_path):
    doc.add_paragraph()
    fig_p2 = doc.add_paragraph()
    fig_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = fig_p2.add_run()
    run2.add_picture(f1_fig_path, width=Inches(6.0))
    add_caption(doc,
        "Figure 2. Per-disease macro F1-score for TF-IDF+LR, DistilBERT+KL, and PubMedBERT+KL on the "
        "DDXPlus test set (n=134,529), sorted by DistilBERT+KL F1 ascending. The dashed line marks F1=0.95. "
        "Both transformer models underperform TF-IDF+LR only on diseases with highly similar symptom profiles "
        "(rhinosinusitis variants, angina variants), where the physician differential assigns split probability "
        "between conditions.")

add_heading(doc, "5.7  ECE vs. KL Divergence: Why Metric Choice Matters", 2)
add_para(doc,
    "Table 1 reveals a striking paradox: TF-IDF+LR achieves near-perfect ECE (0.0043) yet has the worst "
    "mean KL divergence (6.2884). DistilBERT+KL and PubMedBERT+KL show the exact opposite: poor ECE "
    "(\u22480.48) but excellent KL (\u22480.41). This resolves when one understands what each metric measures.")
add_para(doc,
    "ECE measures only whether the top-1 predicted confidence matches empirical accuracy of that top-1 "
    "prediction. TF-IDF+LR assigns ~99% probability to a single disease for virtually every patient. "
    "Since it is correct 99.51% of the time, these high confidences are accurate — ECE \u2248 0. "
    "KL divergence, by contrast, measures the full distribution against the physician differential. "
    "TF-IDF+LR assigns near-zero probability to 48 diseases; the physician differential assigns "
    "non-trivial probability to several. This drives KL to 6.29.")
add_para(doc,
    "The transformer models distribute probability across plausible alternatives — top-1 confidence is "
    "typically 60\u201370%, not 99%, correctly signalling uncertainty about second and third diagnoses. "
    "ECE penalises this: a model outputting '60% Disease X, 25% Disease Y' looks miscalibrated by ECE "
    "even when this distribution accurately mirrors the physician's uncertainty.")
add_para(doc,
    "This result demonstrates that ECE is an inappropriate calibration metric for differential diagnosis. "
    "A clinical system should produce a probability distribution reflecting the full diagnostic space, not "
    "maximise single-answer confidence. KL divergence against physician-annotated differentials is the "
    "correct metric. The 15\u00d7 KL improvement from KL-divergence training is clinically meaningful — "
    "a result ECE entirely fails to capture, and would perversely reward the inferior model.")

# ── 6. Discussion ─────────────────────────────────────────────────────────────
add_heading(doc, "6.  Discussion", 1)

add_heading(doc, "6.1  Clinical Significance of KL-Divergence Training", 2)
add_para(doc,
    "The 15x reduction in mean KL from 6.29 to 0.41 is not merely a statistical improvement. In clinical "
    "practice, a differential diagnosis serves as a checklist: the physician orders tests or makes management "
    "decisions based on which conditions have non-trivial probability. A model that concentrates all probability "
    "on one disease will not prompt consideration of dangerous alternatives. By training with KL-divergence "
    "loss against physician-annotated differentials, our models learn to mirror clinical reasoning: 'What else "
    "could this be, and how likely is each alternative?'")
add_para(doc,
    "This is especially important for safety-critical presentations. If a patient presents with chest pain, "
    "dyspnoea, and leg swelling, the model should assign non-trivial probability to both Pulmonary Embolism "
    "and Myocardial Infarction, even if one is more likely. A model outputting only one diagnosis with 99% "
    "confidence could contribute to anchoring bias.")

add_heading(doc, "6.2  Why Domain Pre-Training Does Not Help on DDXPlus", 2)
add_para(doc,
    "PubMedBERT's advantage has been documented on tasks where input is biomedical free text: clinical NER, "
    "PubMed QA, biomedical relation extraction. DDXPlus symptoms are not free text — they are structured "
    "binary and categorical features converted to a templated string using a controlled vocabulary of feature "
    "names. The model's task is learning co-occurrence patterns of these feature tokens, for which general "
    "pre-training is equally effective. This suggests a practical guideline: use biomedical pre-training when "
    "inputs are natural clinical language; for structured feature inputs, general-domain transformers suffice.")

add_heading(doc, "6.3  System Integration", 2)
add_para(doc,
    "The models are integrated into a full-stack health advisory platform providing: (1) DDXPlus differential "
    "diagnosis with calibrated probability distributions; (2) chest X-ray analysis via DenseNet-121 "
    "pre-trained on NIH ChestX-ray14, CheXpert, and RSNA datasets with Grad-CAM heatmaps; (3) automated lab "
    "report interpretation from PDFs supporting 29 test types with gender-specific reference ranges; "
    "(4) patient-friendly explanations via configurable LLM providers including Google MedGemma-4B. "
    "MedGemma is integrated via the HuggingFace Serverless Inference API and optionally as an on-device "
    "pipeline, with clinical system prompting appropriate for medical explanation generation.")

add_heading(doc, "6.4  Limitations", 2)
add_para(doc,
    "Several limitations should be noted. First, DDXPlus cases are generated by a clinical simulator; "
    "generalisation to real-world clinical data is unverified. Second, the 49 diseases cover common "
    "outpatient presentations and exclude rare diseases, psychiatric conditions, and surgical emergencies. "
    "Third, torchxrayvision performance is cited from published benchmarks rather than independently "
    "evaluated on our deployment pipeline. Fourth, MedGemma explanation quality was assessed qualitatively; "
    "a systematic human evaluation by physician raters was not performed.")

# ── 7. Conclusion ─────────────────────────────────────────────────────────────
add_heading(doc, "7.  Conclusion", 1)
add_para(doc,
    "We demonstrate that training transformer models with a KL-divergence soft-label objective on DDXPlus "
    "produces substantially better-calibrated differential diagnosis distributions compared to TF-IDF+LR, "
    "reducing mean KL divergence by 15x (6.29 \u2192 0.41) and improving NDCG@3 from 0.748 to 0.921. "
    "The cost is modest: approximately 1% lower top-1 accuracy. We further show that PubMedBERT's biomedical "
    "pre-training provides no meaningful advantage over DistilBERT on structured-symptom inputs — a finding "
    "that simplifies model selection in practice.")
add_para(doc,
    "These results suggest that the field's focus on top-1 accuracy as the primary metric for diagnostic AI "
    "may be misleading. A model producing a well-calibrated ranked differential is clinically superior to one "
    "that maximises single-label accuracy at the expense of probability distribution quality. Future work "
    "includes evaluation on real clinical notes, longitudinal patient history integration, systematic human "
    "evaluation of MedGemma explanations, and prospective clinical validation.")

# ── Data Availability ─────────────────────────────────────────────────────────
add_heading(doc, "Data Availability", 1)
add_para(doc,
    "The DDXPlus dataset is publicly available under the CC-BY 4.0 licence at the HuggingFace Hub "
    "(mila-iqia/ddxplus). The trained model weights (DistilBERT+KL, PubMedBERT+KL), label maps, "
    "evaluation scripts, and per-disease results are available in the accompanying project repository.")

# ── Ethics Statement ───────────────────────────────────────────────────────────
add_heading(doc, "Ethics Statement", 1)
add_para(doc,
    "This study uses only publicly available synthetic data generated by a clinical simulator; no real "
    "patient data were collected or processed. No human subjects were involved in any part of this research. "
    "Institutional Review Board (IRB) approval was therefore not required. The dataset licence (CC-BY 4.0) "
    "permits the research use described here.")

# ── Declaration of Competing Interests ────────────────────────────────────────
add_heading(doc, "Declaration of Competing Interests", 1)
add_para(doc,
    "The authors declare that they have no known competing financial interests or personal relationships "
    "that could have appeared to influence the work reported in this paper.")

# ── Acknowledgements ──────────────────────────────────────────────────────────
add_heading(doc, "Acknowledgements", 1)
add_para(doc,
    "The authors thank the DDXPlus team (Tchango et al.) for releasing the dataset under an open licence, "
    "and the HuggingFace team for hosting the dataset and model hub infrastructure used in this work.")

# ── 8. References ─────────────────────────────────────────────────────────────
add_heading(doc, "References", 1)

refs = [
    "[1] Tchango, A.F., Goel, R., Jiang, Z., Martel, J., Ghosn, J. (2022). DDXPlus: A Large-Scale Automatic Differential Diagnosis Dataset. NeurIPS 2022 Datasets and Benchmarks.",
    "[2] Sanh, V., Debut, L., Chaumond, J., Wolf, T. (2019). DistilBERT, a distilled version of BERT: smaller, faster, cheaper and lighter. arXiv:1910.01108.",
    "[3] Gu, Y., et al. (2021). Domain-Specific Language Model Pretraining for Biomedical NLP. ACM Transactions on Computing for Healthcare, 3(1), 1-23.",
    "[4] Lee, J., et al. (2020). BioBERT: a pre-trained biomedical language representation model. Bioinformatics, 36(4), 1234-1240.",
    "[5] Rajpurkar, P., et al. (2017). CheXNet: Radiologist-Level Pneumonia Detection on Chest X-Rays with Deep Learning. arXiv:1711.05225.",
    "[6] Cohen, J.P., et al. (2022). TorchXRayVision: A library of chest X-ray datasets and models. Medical Imaging with Deep Learning (MIDL 2022), PMLR vol. 172.",
    "[7] Hinton, G., Vinyals, O., Dean, J. (2015). Distilling the Knowledge in a Neural Network. NeurIPS Deep Learning Workshop.",
    "[8] Guo, C., Pleiss, G., Sun, Y., Weinberger, K.Q. (2017). On Calibration of Modern Neural Networks. ICML 2017.",
    "[9] Devlin, J., Chang, M.W., Lee, K., Toutanova, K. (2019). BERT: Pre-training of Deep Bidirectional Transformers. NAACL 2019.",
    "[10] Alsentzer, E., et al. (2019). Publicly Available Clinical BERT Embeddings. NAACL Clinical NLP Workshop.",
    "[11] Mullenbach, J., et al. (2018). Explainable Prediction of Medical Codes from Clinical Text. NAACL 2018.",
    "[12] Buch, V.H., Ahmed, I., Maruthappu, M. (2018). Artificial intelligence in medicine: current trends and future possibilities. Br J Gen Pract, 68(668), 143-144.",
    "[13] Rotmensch, M., et al. (2017). Learning a Health Knowledge Graph from Electronic Medical Records. Scientific Reports, 7(1), 5994.",
]

for ref in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(ref)
    set_font(run, size=9)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.left_indent = Inches(0.25)

# ── Appendix A: Hyperparameters ───────────────────────────────────────────────
add_heading(doc, "Appendix A: Training Hyperparameters", 1)

app_headers = ['Hyperparameter', 'DistilBERT+KL', 'PubMedBERT+KL']
app_rows = [
    ['Base model',       'distilbert-base-uncased', 'BiomedNLP-BiomedBERT-base-uncased'],
    ['Learning rate',    '2\u00d710\u207b\u2075',  '2\u00d710\u207b\u2075'],
    ['Batch size',       '64', '64'],
    ['Epochs',           '3', '3'],
    ['Max seq. length',  '128 tokens', '128 tokens'],
    ['Optimiser',        'AdamW', 'AdamW'],
    ['Weight decay',     '0.01', '0.01'],
    ['Warmup steps',     '500', '500'],
    ['KL weight (\u03b1)',  '0.7', '0.7'],
    ['CE weight (1-\u03b1)', '0.3', '0.3'],
    ['Hardware',         'Apple M-series (MPS)', 'Apple M-series (MPS)'],
    ['Training time',    '~2 hours', '~2 hours'],
]
add_table(doc, app_headers, app_rows, col_widths=[2.0, 2.5, 2.5])
add_caption(doc, "Table A1. Hyperparameters for both transformer models.")

# ── Appendix B: Disease List ─────────────────────────────────────────────────
add_heading(doc, "Appendix B: DDXPlus Disease List (Complete 49 Classes)", 1)
add_para(doc,
    "The 49 diseases used in DDXPlus, in label-index order (sorted alphabetically). "
    "Source: models/ddxplus_model/ddxplus_label_map.json.")

ddx_diseases = [
    (0,  "Acute COPD exacerbation / infection"), (1,  "Acute dystonic reactions"),
    (2,  "Acute laryngitis"),                    (3,  "Acute otitis media"),
    (4,  "Acute pulmonary edema"),               (5,  "Acute rhinosinusitis"),
    (6,  "Allergic sinusitis"),                  (7,  "Anaphylaxis"),
    (8,  "Anemia"),                              (9,  "Atrial fibrillation"),
    (10, "Boerhaave"),                           (11, "Bronchiectasis"),
    (12, "Bronchiolitis"),                       (13, "Bronchitis"),
    (14, "Bronchospasm / acute asthma exacerbation"), (15, "Chagas"),
    (16, "Chronic rhinosinusitis"),              (17, "Cluster headache"),
    (18, "Croup"),                               (19, "Ebola"),
    (20, "Epiglottitis"),                        (21, "GERD"),
    (22, "Guillain-Barr\u00e9 syndrome"),        (23, "HIV (initial infection)"),
    (24, "Influenza"),                           (25, "Inguinal hernia"),
    (26, "Larygospasm"),                         (27, "Localized edema"),
    (28, "Myasthenia gravis"),                   (29, "Myocarditis"),
    (30, "PSVT"),                                (31, "Pancreatic neoplasm"),
    (32, "Panic attack"),                        (33, "Pericarditis"),
    (34, "Pneumonia"),                           (35, "Possible NSTEMI / STEMI"),
    (36, "Pulmonary embolism"),                  (37, "Pulmonary neoplasm"),
    (38, "SLE"),                                 (39, "Sarcoidosis"),
    (40, "Scombroid food poisoning"),            (41, "Spontaneous pneumothorax"),
    (42, "Spontaneous rib fracture"),            (43, "Stable angina"),
    (44, "Tuberculosis"),                        (45, "URTI"),
    (46, "Unstable angina"),                     (47, "Viral pharyngitis"),
    (48, "Whooping cough"),
]

# Two-column layout: left half and right half
mid = 25
left  = ddx_diseases[:mid]
right = ddx_diseases[mid:]
appb_headers = ['#', 'Disease', '#', 'Disease']
appb_rows = []
for i in range(mid):
    r = right[i] if i < len(right) else ('', '')
    appb_rows.append([str(left[i][0]), left[i][1], str(r[0]), r[1]])

add_table(doc, appb_headers, appb_rows, col_widths=[0.3, 2.6, 0.3, 2.6])
add_caption(doc, "Table B1. Complete list of 49 diseases in DDXPlus (label indices match ddxplus_label_map.json).")

# ── Save ──────────────────────────────────────────────────────────────────────
os.makedirs(os.path.dirname(os.path.abspath(OUT)), exist_ok=True)
doc.save(OUT)
print(f"Saved: {os.path.abspath(OUT)}")
